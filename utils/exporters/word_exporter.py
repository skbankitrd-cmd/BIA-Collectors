from docx import Document
from docx.shared import Pt, RGBColor
import os
import uuid
from datetime import datetime

class WordExporter:
    """原子化 Word 導出器：全系統統一 Docx 規範"""
    
    def __init__(self):
        self.output_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../../BIA-Web-Solution/server/public/generated_reports"))
        os.makedirs(self.output_dir, exist_ok=True)
        self.primary_color = RGBColor(230, 0, 18) # 台新紅

    def export(self, data: list, report_title: str) -> str:
        """將資料轉為 Word"""
        doc = Document()
        
        # 標題
        heading = doc.add_heading(report_title, 0)
        # 設定標題文字顏色
        run = heading.runs[0]
        run.font.color.rgb = self.primary_color
        
        # 頁首資訊
        p_info = doc.add_paragraph()
        p_info.add_run(f"產出日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        p_info.add_run("BIA 戰略情報系統 - 自動化生成報告").italic = True
        
        doc.add_paragraph("-" * 50)
        
        # 內容列表
        for item in data:
            p = doc.add_paragraph(style='List Bullet')
            # 轉換為字串並美化輸出
            p.add_run(str(item))
            
        file_name = f"DOCX_{report_title}_{uuid.uuid4().hex[:4]}.docx"
        full_path = os.path.join(self.output_dir, file_name)
        doc.save(full_path)
        
        return f"/public/generated_reports/{file_name}"
